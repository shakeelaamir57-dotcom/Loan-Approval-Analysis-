import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
import os

def create_report_docx():
    doc = docx.Document()
    
    # Set page margins (1 inch)
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        
    # Configure default style font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Segoe UI'
    font.size = Pt(11)
    font.color.rgb = RGBColor(0x33, 0x33, 0x33) # Charcoal body text
    
    # Helper to style tables
    def style_table_header(row, color_hex="1E3A8A"):
        for cell in row.cells:
            # Set background color
            shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
            cell._tc.get_or_add_tcPr().append(shading_elm)
            # Set text font, size, and color to white
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.name = 'Segoe UI'
                    run.font.size = Pt(10.5)
                    run.font.bold = True
                    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                    
    def style_table_cells(table):
        for i, row in enumerate(table.rows):
            if i == 0:
                style_table_header(row)
                continue
            # Set background shading for zebra rows
            fill_color = "F1F5F9" if i % 2 == 1 else "FFFFFF"
            for cell in row.cells:
                shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_color}"/>')
                cell._tc.get_or_add_tcPr().append(shading_elm)
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.name = 'Segoe UI'
                        run.font.size = Pt(9.5)
                        
    # Add a custom heading function
    def add_heading_styled(text, level, color_rgb=(0x1E, 0x3A, 0x8A)):
        h = doc.add_heading(level=level)
        run = h.add_run(text)
        run.font.name = 'Segoe UI'
        run.font.color.rgb = RGBColor(*color_rgb)
        run.font.bold = True
        if level == 1:
            run.font.size = Pt(20)
            h.paragraph_format.space_before = Pt(18)
            h.paragraph_format.space_after = Pt(8)
            # Add horizontal line below heading 1
            pBdr = parse_xml(r'<w:pBdr xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
                             r'<w:bottom w:val="single" w:sz="6" w:space="4" w:color="1E3A8A"/>'
                             r'</w:pBdr>')
            h._p.get_or_add_pPr().append(pBdr)
        elif level == 2:
            run.font.size = Pt(15)
            h.paragraph_format.space_before = Pt(12)
            h.paragraph_format.space_after = Pt(6)
        elif level == 3:
            run.font.size = Pt(12)
            h.paragraph_format.space_before = Pt(8)
            h.paragraph_format.space_after = Pt(4)
        return h

    # ----------------------------------------------------
    # PAGE 1: COVER PAGE
    # ----------------------------------------------------
    doc.add_paragraph().paragraph_format.space_before = Pt(120)
    
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_title = p_title.add_run("RETAIL MORTGAGE & CREDIT RISK ANALYTICS")
    run_title.font.size = Pt(28)
    run_title.font.bold = True
    run_title.font.color.rgb = RGBColor(0x1E, 0x3A, 0x8A)
    
    p_subtitle = doc.add_paragraph()
    p_subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_sub = p_subtitle.add_run("End-to-End Loan Approval Prediction & Banking BI Report")
    run_sub.font.size = Pt(14)
    run_sub.font.italic = True
    run_sub.font.color.rgb = RGBColor(0x47, 0x55, 0x69)
    p_subtitle.paragraph_format.space_after = Pt(150)
    
    p_meta = doc.add_paragraph()
    p_meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_meta = p_meta.add_run(
        "Mortgage Lending & Retail Credit Division\n"
        "Analysis using Python (Pandas) · SQL (SQLite/PostgreSQL) · Power BI · SHAP XAI\n\n"
        "Prepared by: Mohammed Aamir Shakeel Ahmad\n"
        "June 2026"
    )
    run_meta.font.size = Pt(11)
    run_meta.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)
    
    doc.add_page_break()
    
    # ----------------------------------------------------
    # PAGE 2: EXECUTIVE SUMMARY & BUSINESS PROBLEM
    # ----------------------------------------------------
    add_heading_styled("1. Executive Summary", level=1)
    p1 = doc.add_paragraph(
        "This report presents a comprehensive risk analytics and machine learning pipeline conducted on housing loan portfolios. "
        "The analysis was executed using Python (Pandas) for data cleaning, SQLite for relational querying, "
        "Scikit-Learn/XGBoost/CatBoost for predictive modeling, SHAP for Explainable AI (XAI) compliance, "
        "and Power BI for executive dashboard visualizations."
    )
    p1.paragraph_format.space_after = Pt(10)
    
    p2 = doc.add_paragraph(
        "Our descriptive analytics revealed a portfolio-level loan approval rate of 68.73%. We identified Credit History "
        "as the single most critical predictor of eligibility, showing a 71.7% approval variance between applicants with "
        "good credit history vs. those without. Furthermore, we determined that household-level financial stability "
        "(defined by combining applicant and co-applicant incomes) significantly reduces default risks compared to assessing solo applicant incomes alone."
    )
    p2.paragraph_format.space_after = Pt(10)
    
    p3 = doc.add_paragraph(
        "We trained and tuned five classification models. The CatBoost Classifier was selected as our production-grade model, "
        "achieving an 85.37% test accuracy and 98.80% sensitivity (recall). To ensure fair lending compliance under the "
        "Equal Credit Opportunity Act, we incorporated SHAP explainability models. SHAP values confirmed that credit decisions "
        "are primarily driven by financial solvency and credit histories, while demographic features (such as gender and marital status) "
        "exhibit near-zero model contribution, assuring compliance with banking fairness guidelines."
    )
    p3.paragraph_format.space_after = Pt(10)
    
    add_heading_styled("2. Business Problem", level=1)
    p4 = doc.add_paragraph(
        "Retail mortgage underwriting involves a balance between credit expansion and asset quality protection. "
        "Manual underwriting processes create operational bottlenecks, increase loan processing time, and introduce subjective bias. "
        "By leveraging machine learning, the bank aims to automate loan pre-approvals while maintaining strict risk controls."
    )
    p4.paragraph_format.space_after = Pt(10)
    
    p5 = doc.add_paragraph("Specifically, this project addresses the following critical banking questions:")
    doc.add_paragraph("• What are the primary demographic and financial drivers of credit defaults?", style='List Bullet')
    doc.add_paragraph("• How much does credit history influence overall mortgage approvals?", style='List Bullet')
    doc.add_paragraph("• Do joint applications (with co-applicants) significantly lower risk profiles?", style='List Bullet')
    doc.add_paragraph("• Can machine learning models predict loan eligibility with high sensitivity (recall) to minimize rejecting creditworthy customers?", style='List Bullet')
    doc.add_paragraph("• How does the bank ensure its automated models comply with fair lending regulations and are free from demographic bias?", style='List Bullet')
    
    doc.add_page_break()

    # ----------------------------------------------------
    # PAGE 3: OBJECTIVES & DATASET DESCRIPTION
    # ----------------------------------------------------
    add_heading_styled("3. Project Objectives", level=1)
    doc.add_paragraph("• Standardize raw datasets by cleaning missing features and handling financial outliers.", style='List Bullet')
    doc.add_paragraph("• Load clean records into an SQL database to run structured portfolio audits.", style='List Bullet')
    doc.add_paragraph("• Perform Exploratory Data Analysis (EDA) to map demographic trends.", style='List Bullet')
    doc.add_paragraph("• Engineer risk-focused financial features (e.g., Loan-to-Income and EMI estimates).", style='List Bullet')
    doc.add_paragraph("• Train and hyperparameter-tune multiple classification algorithms under cross-validation.", style='List Bullet')
    doc.add_paragraph("• Evaluate models using accuracy, precision, recall, and ROC-AUC curves.", style='List Bullet')
    doc.add_paragraph("• Apply SHAP Explainable AI to audit model fairness.", style='List Bullet')
    doc.add_paragraph("• Design a Power BI dashboard template for executive credit risk monitoring.", style='List Bullet')
    
    add_heading_styled("4. Dataset Description", level=1)
    doc.add_paragraph(
        "The project utilizes a housing loan dataset consisting of 614 historical loan files with 13 key categorical and numerical attributes:"
    )
    
    # Create dataset description table
    table_desc = doc.add_table(rows=14, cols=3)
    hdr_cells = table_desc.rows[0].cells
    hdr_cells[0].text = 'Feature Name'
    hdr_cells[1].text = 'Type'
    hdr_cells[2].text = 'Description'
    
    desc_data = [
        ('Loan_ID', 'Categorical', 'Unique loan identifier'),
        ('Gender', 'Categorical', 'Male / Female'),
        ('Married', 'Categorical', 'Applicant marriage status (Yes / No)'),
        ('Dependents', 'Categorical', 'Number of family dependents (0, 1, 2, 3+)'),
        ('Education', 'Categorical', 'Education level (Graduate / Not Graduate)'),
        ('Self_Employed', 'Categorical', 'Self-employment indicator (Yes / No)'),
        ('ApplicantIncome', 'Numerical', 'Monthly income of the primary applicant'),
        ('CoapplicantIncome', 'Numerical', 'Monthly income of the co-applicant'),
        ('LoanAmount', 'Numerical', 'Requested loan amount in thousands ($K)'),
        ('Loan_Amount_Term', 'Numerical', 'Repayment term of the loan in months'),
        ('Credit_History', 'Categorical', 'Credit history meets guidelines (1 = Good, 0 = Bad)'),
        ('Property_Area', 'Categorical', 'Property location (Urban / Semiurban / Rural)'),
        ('Loan_Status', 'Categorical', 'Target Variable: Loan approved (Y / N)')
    ]
    
    for row_idx, data in enumerate(desc_data, start=1):
        row_cells = table_desc.rows[row_idx].cells
        row_cells[0].text = data[0]
        row_cells[1].text = data[1]
        row_cells[2].text = data[2]
        
    style_table_cells(table_desc)
    doc.add_page_break()

    # ----------------------------------------------------
    # PAGE 4: CLEANING PROCESS & SQL DATABASE SETUP
    # ----------------------------------------------------
    add_heading_styled("5. Data Cleaning Process", level=1)
    p_clean = doc.add_paragraph(
        "Before performing modeling, raw data was loaded into a Python environment and preprocessed.\n\n"
        "• Categorical Columns: Missing values in Gender, Married, Dependents, and Self_Employed were imputed with their respective column Mode.\n"
        "• Numerical Columns: Continuous variables like LoanAmount and Loan_Amount_Term were imputed with their column Median to prevent outliers from distorting imputations.\n"
        "• Type Conversions: Credit_History and Loan_Amount_Term were standardized to discrete integers.\n"
        "• Outlier Policy: Outliers in income and loan requests were kept because they represent real high-net-worth applications rather than data errors."
    )
    p_clean.paragraph_format.space_after = Pt(10)
    
    if os.path.exists("Images/outlier_boxplots.png"):
        doc.add_picture("Images/outlier_boxplots.png", width=Inches(5.5))
        p_img = doc.add_paragraph("Figure 1: Boxplot analysis of credit applicant incomes and loan amounts")
        p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_img.runs[0].font.size = Pt(9.5)
        p_img.runs[0].font.italic = True
        
    add_heading_styled("6. SQL Database Ingestion", level=1)
    p_sql = doc.add_paragraph(
        "The preprocessed data was written to a local SQLite database (loan_database.db) to enable direct SQL queries. "
        "We created three structural views to aggregate demographic, geographic, and risk profiles:\n"
        "• v_loan_summary: Evaluates applicant income, loan amounts, and credit history ratios by loan status.\n"
        "• v_risk_profile: Classifies applicants into Low, Medium, and High risk segments.\n"
        "• v_property_analysis: Aggregates volumes, approval rates, and loan averages across property areas."
    )
    p_sql.paragraph_format.space_after = Pt(10)
    doc.add_page_break()

    # ----------------------------------------------------
    # PAGE 5: SQL PORTFOLIO BUSINESS INQUIRIES (PART 1)
    # ----------------------------------------------------
    add_heading_styled("7. SQL Analysis: Business Performance Queries", level=1)
    
    doc.add_paragraph("Q1 Overall Loan Approval Rate — What is the bank's general approval baseline?")
    p_q1_code = doc.add_paragraph("SELECT COUNT(*) AS Total, SUM(CASE WHEN Loan_Status = 'Y' THEN 1 ELSE 0 END) AS Approved, ROUND(SUM(CASE WHEN Loan_Status = 'Y' THEN 1 ELSE 0 END) * 100.0 / COUNT(*), 2) AS Approval_Rate FROM loan_cleaned;")
    p_q1_code.runs[0].font.name = 'Consolas'
    p_q1_code.runs[0].font.size = Pt(9.5)
    doc.add_paragraph("Result: Total: 614, Approved: 422, Approval_Rate: 68.73%\n▶ Business Impact: Establishes the credit baseline. The bank's baseline approval rate is 68.73%, meaning roughly 2 out of 3 applicants receive financing.")
    
    doc.add_paragraph("Q2 Credit History Impact — Does historical repayment compliance predict approval?")
    p_q2_code = doc.add_paragraph("SELECT Credit_History, COUNT(*) AS Total_Applicants, SUM(CASE WHEN Loan_Status = 'Y' THEN 1 ELSE 0 END) AS Approved_Loans, ROUND(SUM(CASE WHEN Loan_Status = 'Y' THEN 1 ELSE 0 END) * 100.0 / COUNT(*), 2) AS Approval_Rate FROM loan_cleaned GROUP BY Credit_History;")
    p_q2_code.runs[0].font.name = 'Consolas'
    p_q2_code.runs[0].font.size = Pt(9.5)
    doc.add_paragraph("Result: Credit_History 0: 7.87% Approval Rate | Credit_History 1: 79.58% Approval Rate\n▶ Business Impact: Confirms credit policy enforcement. Applicants with good credit history have a 79.58% approval rate, whereas those with poor history have only a 7.87% approval rate.")
    
    doc.add_paragraph("Q3 Geographic Performance — Which property areas exhibit the highest credit velocity?")
    p_q3_code = doc.add_paragraph("SELECT Property_Area, COUNT(*) AS Total, ROUND(SUM(CASE WHEN Loan_Status = 'Y' THEN 1 ELSE 0 END) * 100.0 / COUNT(*), 2) AS Approval_Rate FROM loan_cleaned GROUP BY Property_Area;")
    p_q3_code.runs[0].font.name = 'Consolas'
    p_q3_code.runs[0].font.size = Pt(9.5)
    doc.add_paragraph("Result: Rural: 61.45% | Semiurban: 76.82% | Urban: 65.84%\n▶ Business Impact: Semi-urban properties are key drivers, showing a 76.82% approval rate, compared to Urban (65.84%) and Rural (61.45%).")
    
    doc.add_paragraph("Q4 Solo vs. Joint Applications — Do co-applicants improve approval chances?")
    p_q4_code = doc.add_paragraph("SELECT CASE WHEN CoapplicantIncome > 0 THEN 'Joint' ELSE 'Solo' END AS App_Type, ROUND(SUM(CASE WHEN Loan_Status = 'Y' THEN 1 ELSE 0 END) * 100.0 / COUNT(*), 2) AS Approval_Rate FROM loan_cleaned GROUP BY App_Type;")
    p_q4_code.runs[0].font.name = 'Consolas'
    p_q4_code.runs[0].font.size = Pt(9.5)
    doc.add_paragraph("Result: Joint: 72.13% Approval Rate | Solo: 64.98% Approval Rate\n▶ Business Impact: Joint applications have an approval rate of 72.13%, compared to 64.98% for solo applications, supporting the use of combined household income to lower risk.")
    
    doc.add_page_break()

    # ----------------------------------------------------
    # PAGE 6: SQL PORTFOLIO BUSINESS INQUIRIES (PART 2)
    # ----------------------------------------------------
    add_heading_styled("8. SQL Analysis: Advanced Performance Queries", level=1)
    
    doc.add_paragraph("Q5 Portfolio Risk Segmentation — How are applicants distributed across risk tiers?")
    p_q5_code = doc.add_paragraph(
        "WITH RiskCTE AS (\n"
        "    SELECT CASE \n"
        "        WHEN Credit_History = 0 THEN 'High Risk'\n"
        "        WHEN Credit_History = 1 AND (ApplicantIncome+CoapplicantIncome) < 4000 THEN 'Medium Risk'\n"
        "        ELSE 'Low Risk' END AS Risk_Tier, Loan_Status FROM loan_cleaned\n"
        ") SELECT Risk_Tier, COUNT(*) AS Total, ROUND(SUM(CASE WHEN Loan_Status='Y' THEN 1 ELSE 0 END)*100.0/COUNT(*), 2) AS App_Rate FROM RiskCTE GROUP BY Risk_Tier;"
    )
    p_q5_code.runs[0].font.name = 'Consolas'
    p_q5_code.runs[0].font.size = Pt(9.5)
    doc.add_paragraph("Result: Low Risk: 82.35% App_Rate | Medium Risk: 69.57% App_Rate | High Risk: 7.87% App_Rate\n▶ Business Impact: Segmenting applicants by risk helps identify high-quality borrowers. The Low Risk segment has an approval rate of 82.35%, whereas High Risk (poor credit history) has only 7.87%, highlighting where credit policy guidelines are most restrictive.")
    
    doc.add_paragraph("Q6 Income Category Approval Rate — Does applicant income bracket influence approval?")
    p_q6_code = doc.add_paragraph(
        "SELECT CASE \n"
        "    WHEN ApplicantIncome < 3000 THEN 'Low' \n"
        "    WHEN ApplicantIncome BETWEEN 3000 AND 6000 THEN 'Medium' \n"
        "    ELSE 'High' END AS Income_Class, COUNT(*) AS Total, ROUND(SUM(CASE WHEN Loan_Status='Y' THEN 1 ELSE 0 END)*100.0/COUNT(*), 2) AS App_Rate FROM loan_cleaned GROUP BY Income_Class;"
    )
    p_q6_code.runs[0].font.name = 'Consolas'
    p_q6_code.runs[0].font.size = Pt(9.5)
    doc.add_paragraph("Result: Low: 69.3% Approval Rate | Medium: 69.9% Approval Rate | High: 66.7% Approval Rate\n▶ Business Impact: Reveals that raw income brackets have a relatively small baseline impact on approvals, showing that other metrics (like credit history and LTI ratio) carry far more weight in decision-making.")
    
    doc.add_page_break()

    # ----------------------------------------------------
    # PAGE 7: PYTHON EXPLORATORY DATA ANALYSIS (EDA)
    # ----------------------------------------------------
    add_heading_styled("9. Python Exploratory Data Analysis (EDA)", level=1)
    doc.add_paragraph(
        "We conducted Exploratory Data Analysis (EDA) to evaluate distributions, relationships, and correlations.\n\n"
        "• Target Distribution: The class balance of approved vs. rejected loans highlights a typical lending distribution where approved applications represent the majority.\n"
        "• Demographic Breakdown: Checking approvals across demographics showed that education levels and property areas have significant secondary effects on approval rates.\n"
        "• Credit Guidelines: Analysis shows that poor credit history is the strongest negative factor for approval."
    )
    
    # Add EDA Images side by side or vertically
    if os.path.exists("Images/eda_target_distribution.png"):
        doc.add_picture("Images/eda_target_distribution.png", width=Inches(3.8))
    if os.path.exists("Images/eda_categorical_vs_status.png"):
        doc.add_picture("Images/eda_categorical_vs_status.png", width=Inches(3.8))
        
    p_eda_fig = doc.add_paragraph("Figure 2: Portfolio Target Class Distribution and Demographic Breakdown vs. Approval Status")
    p_eda_fig.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_eda_fig.runs[0].font.size = Pt(9.5)
    p_eda_fig.runs[0].font.italic = True
    
    doc.add_page_break()

    # ----------------------------------------------------
    # PAGE 8: FEATURE ENGINEERING
    # ----------------------------------------------------
    add_heading_styled("10. Feature Engineering", level=1)
    doc.add_paragraph(
        "To capture repayment capacity and applicant risk profiles, we engineered several financial features:\n\n"
        "1. TotalIncome: Combines applicant and co-applicant incomes to measure household purchasing power.\n"
        "2. Loan-to-Income Ratio (LTI %): Measures the scale of requested credit exposure relative to monthly household income.\n"
        "3. Estimated EMI: Estimates the monthly debt servicing cost, assuming a simple interest structure.\n"
        "4. Risk Category: A rule-based indicator classifying applicants as Low, Medium, or High Risk using credit history and total income thresholds."
    )
    
    if os.path.exists("Images/eda_correlation_heatmap.png"):
        doc.add_picture("Images/eda_correlation_heatmap.png", width=Inches(4.5))
        p_heat_fig = doc.add_paragraph("Figure 3: Correlation heatmap of numerical features and engineered features with target status")
        p_heat_fig.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_heat_fig.runs[0].font.size = Pt(9.5)
        p_heat_fig.runs[0].font.italic = True
        
    doc.add_page_break()

    # ----------------------------------------------------
    # PAGE 9: MACHINE LEARNING MODEL BUILDING & COMPARISON
    # ----------------------------------------------------
    add_heading_styled("11. Machine Learning Model Comparisons", level=1)
    doc.add_paragraph(
        "We implemented a machine learning pipeline using ColumnTransformer to scale numerical features and one-hot encode categorical features. "
        "Models were trained on an 80-20 stratified split and evaluated using 5-Fold Cross Validation. "
        "We performed hyperparameter tuning via GridSearchCV on Random Forest, XGBoost, and CatBoost."
    )
    
    # Model comparison table
    table_ml = doc.add_table(rows=6, cols=6)
    hdr_ml = table_ml.rows[0].cells
    hdr_ml[0].text = 'Model Pipeline'
    hdr_ml[1].text = 'Accuracy'
    hdr_ml[2].text = 'Precision'
    hdr_ml[3].text = 'Recall'
    hdr_ml[4].text = 'F1-Score'
    hdr_ml[5].text = 'ROC-AUC'
    
    ml_data = [
        ('CatBoost Classifier (Tuned)', '85.37%', '82.80%', '98.80%', '90.10%', '0.842'),
        ('XGBoost Classifier (Tuned)', '83.74%', '82.47%', '96.39%', '88.89%', '0.829'),
        ('Random Forest (Tuned)', '82.93%', '81.63%', '96.39%', '88.39%', '0.821'),
        ('Logistic Regression (Baseline)', '81.30%', '79.21%', '96.39%', '86.96%', '0.803'),
        ('Decision Tree (Baseline)', '72.36%', '79.27%', '80.24%', '79.75%', '0.681')
    ]
    
    for row_idx, data in enumerate(ml_data, start=1):
        row_cells = table_ml.rows[row_idx].cells
        for col_idx, val in enumerate(data):
            row_cells[col_idx].text = val
            
    style_table_cells(table_ml)
    
    if os.path.exists("Images/evaluation_roc_curves.png"):
        doc.add_picture("Images/evaluation_roc_curves.png", width=Inches(3.8))
        p_roc_fig = doc.add_paragraph("Figure 4: Receiver Operating Characteristic (ROC) curves comparison")
        p_roc_fig.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_roc_fig.runs[0].font.size = Pt(9.5)
        p_roc_fig.runs[0].font.italic = True
        
    doc.add_page_break()

    # ----------------------------------------------------
    # PAGE 10: MODEL PREDICTIONS ON TEST SET & SHAP EXPLAINABILITY
    # ----------------------------------------------------
    add_heading_styled("12. Model Predictions on Test Set Data", level=1)
    doc.add_paragraph(
        "Below is a sample of 20 test set applications processed by our production CatBoost pipeline, showing actual statuses, predicted probability scores, and final predictions:"
    )
    
    # Table of test predictions
    table_test = doc.add_table(rows=21, cols=9)
    hdr_test = table_test.rows[0].cells
    hdr_test[0].text = 'Row'
    hdr_test[1].text = 'Income'
    hdr_test[2].text = 'Loan ($K)'
    hdr_test[3].text = 'Credit'
    hdr_test[4].text = 'Area'
    hdr_test[5].text = 'Actual'
    hdr_test[6].text = 'Prob'
    hdr_test[7].text = 'Pred'
    hdr_test[8].text = 'Audit'
    
    test_rows_data = [
        ('150', '6,277', '118', '0', 'Rural', 'Rejected', '0.1669', 'Rejected', 'Correct'),
        ('559', '6,486', '182', '1', 'Semiurban', 'Approved', '0.6952', 'Approved', 'Correct'),
        ('598', '9,963', '180', '1', 'Rural', 'Approved', '0.6626', 'Approved', 'Correct'),
        ('235', '6,760', '170', '1', 'Rural', 'Approved', '0.6827', 'Approved', 'Correct'),
        ('145', '6,816', '100', '1', 'Semiurban', 'Approved', '0.7128', 'Approved', 'Correct'),
        ('191', '12,000', '164', '1', 'Semiurban', 'Rejected', '0.6740', 'Approved', 'False Pos'),
        ('557', '10,139', '260', '1', 'Semiurban', 'Approved', '0.7003', 'Approved', 'Correct'),
        ('470', '5,529', '162', '1', 'Semiurban', 'Approved', '0.7055', 'Approved', 'Correct'),
        ('88', '8,566', '210', '1', 'Urban', 'Approved', '0.6584', 'Approved', 'Correct'),
        ('386', '3,946', '132', '1', 'Semiurban', 'Approved', '0.7276', 'Approved', 'Correct'),
        ('380', '5,833', '128', '1', 'Semiurban', 'Approved', '0.7312', 'Approved', 'Correct'),
        ('335', '9,993', '70', '1', 'Semiurban', 'Approved', '0.6988', 'Approved', 'Correct'),
        ('368', '6,325', '175', '1', 'Semiurban', 'Approved', '0.7027', 'Approved', 'Correct'),
        ('60', '6,296', '120', '1', 'Urban', 'Approved', '0.7017', 'Approved', 'Correct'),
        ('569', '5,230', '104', '0', 'Urban', 'Rejected', '0.1686', 'Rejected', 'Correct'),
        ('517', '4,874', '123', '0', 'Semiurban', 'Rejected', '0.1877', 'Rejected', 'Correct'),
        ('500', '4,328', '113', '1', 'Rural', 'Approved', '0.6834', 'Approved', 'Correct'),
        ('399', '3,300', '103', '0', 'Semiurban', 'Rejected', '0.1991', 'Rejected', 'Correct'),
        ('414', '5,386', '178', '0', 'Semiurban', 'Rejected', '0.2417', 'Rejected', 'Correct'),
        ('508', '5,492', '188', '1', 'Urban', 'Approved', '0.6774', 'Approved', 'Correct')
    ]
    
    for row_idx, data in enumerate(test_rows_data, start=1):
        row_cells = table_test.rows[row_idx].cells
        for col_idx, val in enumerate(data):
            row_cells[col_idx].text = val
            
    style_table_cells(table_test)
    
    add_heading_styled("13. Model Interpretability & Compliance (SHAP XAI)", level=1)
    if os.path.exists("Images/shap_beeswarm_plot.png"):
        doc.add_picture("Images/shap_beeswarm_plot.png", width=Inches(3.8))
        p_shap_fig = doc.add_paragraph("Figure 5: SHAP Beeswarm plot ranking features by predictive importance")
        p_shap_fig.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_shap_fig.runs[0].font.size = Pt(9.5)
        p_shap_fig.runs[0].font.italic = True
        
    doc.add_page_break()

    # ----------------------------------------------------
    # PAGE 11: POWER BI SUMMARY & RECOMMENDATIONS
    # ----------------------------------------------------
    add_heading_styled("14. Power BI Dashboard Layout & Business Recommendations", level=1)
    
    if os.path.exists("Images/loan_dashboard_mockup.jpg"):
        doc.add_picture("Images/loan_dashboard_mockup.jpg", width=Inches(5.5))
        p_dash_fig = doc.add_paragraph("Figure 6: High-fidelity layout of the Power BI mortgage analytics dashboard")
        p_dash_fig.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_dash_fig.runs[0].font.size = Pt(9.5)
        p_dash_fig.runs[0].font.italic = True
        
    add_heading_styled("14.1 Policy Recommendations for Credit Risk Teams:", level=2)
    doc.add_paragraph(
        "1. Mandate Household Income Reporting: Require both applicant and co-applicant incomes to support mortgage files, lowering overall credit default risk.\n"
        "2. Target Semi-Urban Portfolios: Focus mortgage expansion campaigns in semi-urban locations showing strong 76.82% approval rates and low relative defaults.\n"
        "3. Alternative Credit Scorecards: Implement utility bill and transaction-based scoring engines for applicants with Credit_History = 0 to capture creditworthy first-time buyers."
    )
    
    doc.save("banking_analytics_performance_report.docx")
    print("Report 'banking_analytics_performance_report.docx' created successfully in docx format!")

if __name__ == "__main__":
    create_report_docx()
